from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "presentation-offre-prospection-olithea-notebooklm.md"
OUTPUT = BASE_DIR / "presentation-offre-prospection-olithea-notebooklm.docx"

OLIVE = "647A0B"
DARK_GREEN = "2F3B24"
BROWN = "854F38"
CREAM = "F7F9EC"
LIGHT_GREEN = "EEF3DD"
TEXT = "374151"
LIGHT_GRAY = "F3F4F6"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_bottom_border(paragraph, color: str = OLIVE, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_inline_runs(paragraph, text: str, color: str | None = None) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BROWN)
        if color and not token.startswith("`"):
            run.font.color.rgb = RGBColor.from_string(color)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def configure_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    title.paragraph_format.space_after = Pt(14)

    heading_specs = {
        "Heading 1": (21, DARK_GREEN, 18, 8),
        "Heading 2": (16, OLIVE, 14, 6),
        "Heading 3": (13, BROWN, 10, 4),
        "Heading 4": (11, DARK_GREEN, 8, 3),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    quote_style = document.styles["Quote"]
    quote_style.font.name = "Aptos"
    quote_style.font.size = Pt(11)
    quote_style.font.italic = False
    quote_style.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    quote_style.paragraph_format.left_indent = Cm(0.7)
    quote_style.paragraph_format.right_indent = Cm(0.5)
    quote_style.paragraph_format.space_before = Pt(5)
    quote_style.paragraph_format.space_after = Pt(8)

    if "Code Block" not in [style.name for style in document.styles]:
        code_style = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(8.5)
        code_style.font.color.rgb = RGBColor.from_string(TEXT)
        code_style.paragraph_format.left_indent = Cm(0.4)
        code_style.paragraph_format.right_indent = Cm(0.4)
        code_style.paragraph_format.space_after = Pt(0)

    header = section.header.paragraphs[0]
    header.text = "OLITHEA  |  Une offre, deux parcours"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(OLIVE)
    set_paragraph_bottom_border(header, color="D9E2B8", size="4")

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    footer.runs[0].font.name = "Aptos"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("6B7280")

    return document


def parse_table(lines: list[str], start: int, document: Document) -> int:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
        rows.append(values)
        index += 1

    if len(rows) < 2:
        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, lines[start])
        return start + 1

    separator = rows[1]
    is_separator = all(re.fullmatch(r":?-{3,}:?", value.replace(" ", "")) for value in separator)
    if not is_separator:
        for row in rows:
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, " | ".join(row))
        return index

    data_rows = [rows[0], *rows[2:]]
    column_count = max(len(row) for row in data_rows)
    table = document.add_table(rows=len(data_rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row_values in enumerate(data_rows):
        for col_index in range(column_count):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            value = row_values[col_index] if col_index < len(row_values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, value, WHITE if row_index == 0 else TEXT)
            for run in paragraph.runs:
                run.font.size = Pt(8.8)
                if row_index == 0:
                    run.bold = True
            set_cell_shading(cell, OLIVE if row_index == 0 else (CREAM if row_index % 2 == 0 else WHITE))
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return index


def render_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    in_code = False
    first_h1 = True

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            index += 1
            continue

        if in_code:
            paragraph = document.add_paragraph(style="Code Block")
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(line)
            p_pr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_GRAY)
            p_pr.append(shd)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            index = parse_table(lines, index, document)
            continue

        if stripped == "---":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            set_paragraph_bottom_border(paragraph, color="D9E2B8", size="6")
            index += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1 and first_h1:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_runs(paragraph, text)
                first_h1 = False
            else:
                if text.startswith("0. Mode d'emploi") or (level == 1 and text.startswith("PARTIE")):
                    document.add_page_break()
                paragraph = document.add_paragraph(style=f"Heading {level}")
                add_inline_runs(paragraph, text)
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = document.add_paragraph(style="Quote")
            add_inline_runs(paragraph, stripped.lstrip("> "))
            set_paragraph_bottom_border(paragraph, color=OLIVE, size="10")
            index += 1
            continue

        bullet_match = re.match(r"^\s*-\s+(.*)$", line)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline_runs(paragraph, bullet_match.group(1))
            index += 1
            continue

        number_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline_runs(paragraph, number_match.group(1))
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, stripped.rstrip("  "))
        index += 1


def add_cover_accents(document: Document) -> None:
    paragraphs = document.paragraphs
    if not paragraphs:
        return
    title = paragraphs[0]
    title.paragraph_format.space_before = Pt(55)
    set_paragraph_bottom_border(title, color=OLIVE, size="14")

    for paragraph in paragraphs[1:8]:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(DARK_GREEN)


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    document = configure_document()
    document.core_properties.title = "Une offre, deux parcours : de la découverte au rendez-vous"
    document.core_properties.subject = "Dossier de présentation Olithea pour NotebookLM"
    document.core_properties.author = "Olithea"
    document.core_properties.keywords = "Olithea, offre, prospection, praticiens, tunnel de vente, newsletter"
    render_markdown(document, markdown)
    add_cover_accents(document)
    document.save(OUTPUT)
    print(f"Created {OUTPUT} ({OUTPUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
